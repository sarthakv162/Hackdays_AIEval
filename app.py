import streamlit as st
import os
import re
import pandas as pd
from PIL import Image
import pytesseract
from docx import Document
import pdfplumber
import google.generativeai as genai

#Setup
st.set_page_config(page_title="AIEval - DTU Assignment Evaluator", layout="wide")
st.title(" AIEval - DTU Assignment Evaluator")
st.caption("Upload assignment submission and answer key to auto-evaluate using Gemini")

tab1, tab2 = st.tabs([" Evaluate", " Dashboard"])
results = []

def extract_text_from_file(file_path):
    if file_path.name.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    else:
        raise ValueError("Unsupported file type. Please upload a .docx or .pdf")

def extract_images_from_docx(file_obj):
    temp_path = "temp.docx"
    with open(temp_path, "wb") as f:
        f.write(file_obj.getbuffer())
    doc = Document(temp_path)
    os.makedirs("images", exist_ok=True)
    rels = doc.part._rels
    paths = []
    for i, rel in enumerate(rels.values(), 1):
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            path = f"images/image_{i}.png"
            with open(path, "wb") as f:
                f.write(image_data)
            paths.append(path)
    return paths

def extract_text_from_images(paths):
    ocr = {}
    for p in paths:
        try:
            img = Image.open(p)
            ocr[p] = pytesseract.image_to_string(img)
        except Exception as e:
            ocr[p] = f"Error: {e}"
    return ocr

def segment_by_questions(text):
    pattern = r"(?:^|\n)(?:Q)?(\d{1,2})[).:]"
    parts = re.split(pattern, text)
    qmap = {}
    for i in range(1, len(parts), 2):
        qmap[parts[i].strip()] = parts[i+1].strip()
    return qmap

def extract_student_name(text):
    prompt = f"""
You are an intelligent assistant helping extract student details.

From the following assignment submission content, extract only the **student's full name**. If no clear name is found, respond with "Anonymous".

Text:
{text[:1500]}  # Limiting to first 1500 characters to save token budget.

Respond in format:
Name: <full name or Anonymous>
"""
    try:
        response = model.generate_content(prompt)
        lines = response.text.strip().splitlines()
        for line in lines:
            if line.lower().startswith("name:"):
                return line.split(":", 1)[1].strip()
        return "Anonymous"
    except:
        return "Anonymous"

def score_answer_with_gemini(q_num, q_text, model_ans, student_ans):
    prompt = f"""
You are an experienced university examiner at DTU, evaluating technical assignment answers.

Assess the student's answer by comparing it to the model answer provided by the professor. I want you to be strict, harsh cut marks wherever can

Evaluation Criteria:
- Conceptual correctness (Has the core logic or explanation been captured?)
- Completeness (Are most key points covered?)
- Relevance (Does the answer stay on-topic and focused?)
- Terminology and structure (Is the language appropriate for a university-level answer?)

Be firm but fair:
- Award partial marks for partially correct but relevant attempts.
- Do not reward vague, generic, or irrelevant responses.
- Slight leniency is allowed if the student shows understanding, even if the wording is different.

---

Question {q_num}: {q_text}

 Model Answer:
{model_ans}

 Student Answer:
{student_ans}

---

Return only:
Score: X/10  
Feedback: One clear, academic sentence justifying the score. Keep it objective and constructive.
"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {e}"

def detect_ai_generated_answer(student_answer, question_text):
    prompt = f"""
You are an AI text detection expert. Think wisely before you mark it AI-generated or human-written

Based on the style, structure, and tone of the following student answer, judge whether it was likely written by a large language model (like ChatGPT or Gemini) or a human.

Question: {question_text}

Answer:
{student_answer}

Respond with one of the following:
- Likely AI-generated
- Likely human-written
- Uncertain

Also provide a short justification.

Format:
Verdict: <one of the above>
Reason: <brief explanation>
"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {e}"

def adjust_score_for_ai(score_text, ai_verdict, penalty=5):
    match = re.search(r"Score:\s*(\d+(?:\.\d+)?)/10", score_text)
    if not match:
        return score_text

    score = float(match.group(1))

    verdict_match = re.search(r"Verdict:\s*(Likely AI-generated)", ai_verdict, re.IGNORECASE)
    is_ai_generated = bool(verdict_match)

    adjusted = max(score - penalty, 0) if is_ai_generated else score
    result = re.sub(r"Score:\s*\d+(?:\.\d+)?/10", f"Score: {adjusted}/10", score_text)
    if adjusted != score:
        result += f"\n(Note: -{penalty} penalty applied due to AI-generated suspicion.)"
    return result

# Evaluate Tab 
with tab1:
    student_file = st.file_uploader("Upload Student Assignment (.docx or .pdf)", type=["pdf", "docx"])
    answer_key_file = st.file_uploader("Upload Answer Key (.docx or .pdf)", type=["pdf", "docx"])

    if student_file and answer_key_file:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel(model_name="models/gemini-2.5-flash")

        with st.spinner("Processing files..."):
            text = extract_text_from_file(student_file)
            ocr = ""
            if student_file.name.endswith(".docx"):
                imgs = extract_images_from_docx(student_file)
                ocr_texts = extract_text_from_images(imgs)
                ocr = "\n".join(ocr_texts.values())
            full_text = text + "\n" + ocr
            questions = segment_by_questions(full_text)

            student_name = extract_student_name(full_text)

            key_text = extract_text_from_file(answer_key_file)
            if answer_key_file.name.endswith(".docx"):
                key_imgs = extract_images_from_docx(answer_key_file)
                key_ocr = extract_text_from_images(key_imgs)
                key_text += "\n" + "\n".join(key_ocr.values())
            model_answers = segment_by_questions(key_text)

            expected = set(model_answers.keys())
            attempted = set(questions.keys())
            missing = expected - attempted

            if missing:
                st.warning("Unattempted Questions: " + ", ".join(sorted(missing)))

            for q_num, ans in questions.items():
                if not ans.strip() or len(ans.split()) < 10:
                    st.info(f" Q{q_num} skipped: too short.")
                    continue
                model_ans = model_answers.get(q_num)
                if not model_ans:
                    st.warning(f" Q{q_num} skipped: no model answer.")
                    continue

                score_output = score_answer_with_gemini(q_num, f"Q{q_num}", model_ans, ans)
                ai_check = detect_ai_generated_answer(ans, f"Q{q_num}")
                adjusted_score = adjust_score_for_ai(score_output, ai_check)

                st.markdown(f"---\n### Q{q_num}")
                st.markdown(f"**Evaluation:**\n\n{adjusted_score.strip()}")
                st.markdown(f"**AI Detection:**\n\n{ai_check.strip()}")

                results.append({
                    "Student": student_name,
                    "Question": f"Q{q_num}",
                    "Score": adjusted_score.split('\n')[0].replace("Score: ", "").replace("/10", ""),
                    "AI Verdict": ai_check.split('\n')[0].replace("Verdict: ", "")
                })

#  Dashboard Tab 
with tab2:
    st.subheader(" Evaluated Results Dashboard")
    if results:
        df = pd.DataFrame(results)
        st.dataframe(df)
        csv = df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label=" Download CSV",
            data=csv,
            file_name="student_evaluations.csv",
            mime="text/csv"
        )
    else:
        st.info("No evaluations yet. Submit an assignment to begin.")


